import cv2
import pytesseract
import numpy as np
import json
import os
from pdf2image import convert_from_path
from sklearn.cluster import KMeans
from docx import Document
from docx.shared import Inches

# Specify Tesseract executable path
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Input and output directories
INPUT_DIR = r"D:\google\thp1\data"
OUTPUT_DIR = r"D:\google\thp1\output"

def preprocess_image(image):
    # Convert to grayscale and apply adaptive thresholding
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 2)
    # Denoise the thresholded image
    denoised = cv2.fastNlMeansDenoising(thresh)
    return denoised, gray

def detect_lines(image):
    # Preprocess image and detect horizontal and vertical lines
    thresh, _ = preprocess_image(image)
    horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
    vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
    horizontal_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
    vertical_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
    return horizontal_lines, vertical_lines

def estimate_spacing(coords):
    # Estimate row spacing using clustering if enough coordinates exist
    if len(coords) < 2:
        return 40
    diffs = np.diff(sorted(coords))
    diffs = diffs.reshape(-1, 1)
    kmeans = KMeans(n_clusters=2, random_state=0).fit(diffs)
    centers = sorted(kmeans.cluster_centers_.flatten())
    return centers[1], centers[0]

def get_row_coordinates(horizontal_lines, image_height):
    # Extract y-coordinates of horizontal lines
    h_contours, _ = cv2.findContours(horizontal_lines, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    h_coords = sorted([cv2.boundingRect(cnt)[1] for cnt in h_contours])
    h_coords = [0] + h_coords + [image_height]
    return h_coords

def get_col_coordinates(vertical_lines, image_width):
    # Extract x-coordinates of vertical lines
    v_contours, _ = cv2.findContours(vertical_lines, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    v_coords = sorted([cv2.boundingRect(cnt)[0] for cnt in v_contours])
    v_coords = [0] + v_coords + [image_width]
    return v_coords

def extract_text_in_cell(image, x, y, w, h, lang='chi_sim+eng'):
    # Extract text from a cell region using OCR
    if w <= 0 or h <= 0:
        return []
    roi = image[y:y + h, x:x + w]
    if roi.size == 0:
        return []
    roi_gray = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY)
    _, roi_thresh = cv2.threshold(roi_gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    d = pytesseract.image_to_data(roi_thresh, lang=lang, output_type=pytesseract.Output.DICT)
    blocks = []
    for i in range(len(d['text'])):
        if d['text'][i].strip():
            blocks.append({
                'text': d['text'][i].strip(),
                'x': d['left'][i],
                'y': d['top'][i],
                'w': d['width'][i],
                'h': d['height'][i],
                'center_y': d['top'][i] + d['height'][i] / 2
            })
    return blocks

def process_multi_line_cell(cell_blocks, small_spacing):
    # Process multi-line text within a cell
    if not cell_blocks:
        return []
    if len(cell_blocks) == 1:
        return [{'sub_row': 1, 'content': cell_blocks[0]['text']}]
    sorted_blocks = sorted(cell_blocks, key=lambda b: b['y'])
    sub_rows = []
    sub_row_idx = 1
    current_sub_row = [sorted_blocks[0]]
    for block in sorted_blocks[1:]:
        if (block['y'] - current_sub_row[-1]['y']) < small_spacing:
            current_sub_row.append(block)
        else:
            content = ' '.join(b['text'] for b in sorted(current_sub_row, key=lambda b: b['y']))
            sub_rows.append({'sub_row': sub_row_idx, 'content': content})
            sub_row_idx += 1
            current_sub_row = [block]
    content = ' '.join(b['text'] for b in sorted(current_sub_row, key=lambda b: b['y']))
    sub_rows.append({'sub_row': sub_row_idx, 'content': content})
    return sub_rows

def extract_form_data(image, lang='chi_sim+eng'):
    # Extract form data from an image
    horizontal_lines, vertical_lines = detect_lines(image)
    h_coords = get_row_coordinates(horizontal_lines, image.shape[0])
    v_coords = get_col_coordinates(vertical_lines, image.shape[1])
    row_spacing, small_spacing = estimate_spacing(h_coords)
    data = []
    for row_idx in range(len(h_coords) - 1):
        y1, y2 = h_coords[row_idx], h_coords[row_idx + 1]
        for col_idx in range(len(v_coords) - 1):
            x1, x2 = v_coords[col_idx], v_coords[col_idx + 1]
            cell_blocks = extract_text_in_cell(image, x1, y1, x2 - x1, y2 - y1, lang)
            sub_rows = process_multi_line_cell(cell_blocks, small_spacing)
            for sub_row in sub_rows:
                entry = {
                    "row": row_idx + 1,
                    "col": col_idx + 1,
                    "content": sub_row['content']
                }
                if len(sub_rows) > 1:
                    entry['sub_row'] = sub_row['sub_row']
                data.append(entry)
    return data

def process_pdf(input_path, output_path, lang='chi_sim+eng'):
    # Process a PDF file to extract form data and save as JSON
    try:
        images = convert_from_path(input_path)
        all_data = []
        for page_num, image in enumerate(images, 1):
            input_img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            page_data = extract_form_data(input_img, lang)
            all_data.extend(page_data)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(all_data, f, indent=4, ensure_ascii=False)
        print(f"Processed: {input_path} -> {output_path}")
    except Exception as e:
        print(f"Error processing: {input_path}, Error: {str(e)}")

def json_to_word(json_path, output_docx_path):
    # Convert JSON data to a Word document with tables
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            cells = json.load(f)

        if not cells:
            print(f"JSON file {json_path} is empty, skipping Word generation")
            return

        doc = Document()

        if 'm1' in json_path.lower():
            # m1.pdf: Metadata and data table
            metadata_cells = [c for c in cells if c['row'] < 10]
            data_cells = [c for c in cells if c['row'] >= 10]

            # Metadata table
            if metadata_cells:
                metadata_rows = sorted(set(c['row'] for c in metadata_cells))
                max_meta_col = 1  # Single-column layout
                meta_table = doc.add_table(rows=len(metadata_rows), cols=max_meta_col)
                meta_table.style = 'Table Grid'

                for idx, row in enumerate(metadata_rows):
                    row_cells = sorted([c for c in metadata_cells if c['row'] == row], key=lambda c: c['col'])
                    content = ' '.join(
                        c['content'].replace('|', '').replace('“', '').replace('”', '').replace('!', '').strip() for c
                        in row_cells)
                    meta_table.rows[idx].cells[0].text = content
                    print(f"Metadata row: row={row}, content={content}")

                doc.add_paragraph()

            # Data table
            if data_cells:
                row_values = sorted(set(c['row'] for c in data_cells))
                row_map = {row: idx for idx, row in enumerate(row_values)}
                col_map = {2: 0, 3: 1, 4: 2, 5: 3, 7: 4, 9: 5, 11: 6, 12: 7}  # Fixed mapping
                max_row = len(row_values)
                max_col = 8

                # Build table content matrix
                table_data = [['' for _ in range(max_col)] for _ in range(max_row)]
                occupied = set()

                for cell in sorted(data_cells, key=lambda c: (c['row'], c['col'], c.get('sub_row', 1))):
                    orig_row = cell['row']
                    orig_col = cell['col']
                    content = cell['content'].replace('|', '').replace('“', '').replace('”', '').replace('—~CO—™”W',
                                                                                                         '').replace(
                        'at ee', '').strip()
                    is_part_of_cell = 'sub_row' in cell

                    row_idx = row_map.get(orig_row, -1)
                    col_idx = col_map.get(orig_col, -1)

                    if row_idx == -1 or col_idx == -1:
                        print(f"Skipping invalid cell: row={orig_row}, col={orig_col}, content={content}")
                        continue

                    if (row_idx, col_idx) in occupied:
                        continue

                    if is_part_of_cell:
                        sub_cells = [c for c in data_cells if
                                     c['row'] == orig_row and c['col'] == orig_col and 'sub_row' in c]
                        sub_cells = sorted(sub_cells, key=lambda c: c['sub_row'])
                        if sub_cells:
                            content = '\n'.join(
                                c['content'].replace('|', '').replace('“', '').replace('”', '').strip() for c in
                                sub_cells)
                            for sub_cell in sub_cells:
                                sub_row_idx = row_map.get(sub_cell['row'], -1)
                                sub_col_idx = col_map.get(sub_cell['col'], -1)
                                if sub_row_idx != -1 and sub_col_idx != -1:
                                    occupied.add((sub_row_idx, sub_col_idx))

                    table_data[row_idx][col_idx] = content
                    occupied.add((row_idx, col_idx))
                    print(f"Data table cell: row={row_idx}, col={col_idx}, content={content}")

                # Create and populate table
                table = doc.add_table(rows=max_row, cols=max_col)
                table.style = 'Table Grid'

                for row_idx in range(max_row):
                    for col_idx in range(max_col):
                        cell_content = table_data[row_idx][col_idx]
                        target_cell = table.rows[row_idx].cells[col_idx]
                        colspan = 1
                        if cell_content.lower() in ['not examined', 'dmso'] or len(cell_content) > 15:
                            colspan = min(3, max_col - col_idx)
                            for i in range(1, colspan):
                                if col_idx + i < max_col and (row_idx, col_idx + i) not in occupied:
                                    target_cell.merge(table.rows[row_idx].cells[col_idx + i])
                                    occupied.add((row_idx, col_idx + i))
                                    table_data[row_idx][col_idx + i] = ''  # Clear merged cell content
                        target_cell.text = cell_content

        else:
            # m0.pdf: Single table
            row_values = sorted(set(c['row'] for c in cells))
            row_map = {row: idx for idx, row in enumerate(row_values)}
            col_map = {2: 0, 3: 1, 4: 2, 7: 3, 8: 4, 10: 5, 11: 6, 12: 7, 13: 8, 17: 9}  # Fixed mapping

            max_row = len(row_values)
            max_col = 10

            # Build table content matrix
            table_data = [['' for _ in range(max_col)] for _ in range(max_row)]
            occupied = set()

            for cell in sorted(cells, key=lambda c: (c['row'], c['col'], c.get('sub_row', 1))):
                orig_row = cell['row']
                orig_col = cell['col']
                content = cell['content'].replace('|', '').replace('“', '').replace('”', '').strip()
                is_part_of_cell = 'sub_row' in cell

                row_idx = row_map.get(orig_row, -1)
                col_idx = col_map.get(orig_col, -1)

                if row_idx == -1 or col_idx == -1:
                    print(f"Skipping invalid cell: row={orig_row}, col={orig_col}, content={content}")
                    continue

                if (row_idx, col_idx) in occupied:
                    continue

                if is_part_of_cell:
                    sub_cells = [c for c in cells if c['row'] == orig_row and c['col'] == orig_col and 'sub_row' in c]
                    sub_cells = sorted(sub_cells, key=lambda c: c['sub_row'])
                    if sub_cells:
                        content = '\n'.join(
                            c['content'].replace('|', '').replace('“', '').replace('”', '').strip() for c in sub_cells)
                        for sub_cell in sub_cells:
                            sub_row_idx = row_map.get(sub_cell['row'], -1)
                            sub_col_idx = col_map.get(sub_cell['col'], -1)
                            if sub_row_idx != -1 and sub_col_idx != -1:
                                occupied.add((sub_row_idx, sub_col_idx))

                table_data[row_idx][col_idx] = content
                occupied.add((row_idx, col_idx))
                print(f"Table cell: row={row_idx}, col={col_idx}, content={content}")

            # Create and populate table
            table = doc.add_table(rows=max_row, cols=max_col)
            table.style = 'Table Grid'

            for row_idx in range(max_row):
                for col_idx in range(max_col):
                    cell_content = table_data[row_idx][col_idx]
                    target_cell = table.rows[row_idx].cells[col_idx]
                    colspan = 1
                    if len(cell_content) > 20:
                        colspan = min(2, max_col - col_idx)
                        for i in range(1, colspan):
                            if col_idx + i < max_col and (row_idx, col_idx + i) not in occupied:
                                target_cell.merge(table.rows[row_idx].cells[col_idx + i])
                                occupied.add((row_idx, col_idx + i))
                                table_data[row_idx][col_idx + i] = ''
                    target_cell.text = cell_content

        doc.save(output_docx_path)
        print(f"Word document generated: {json_path} -> {output_docx_path}")
    except Exception as e:
        print(f"Error generating Word document: {json_path}, Error: {str(e)}")

def main():
    # Create output directory if it doesn't exist
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    # Process all PDF files in the input directory
    for filename in os.listdir(INPUT_DIR):
        if filename.lower().endswith('.pdf'):
            input_path = os.path.join(INPUT_DIR, filename)
            base_name = os.path.splitext(filename)[0]
            output_json_path = os.path.join(OUTPUT_DIR, f"{base_name}.json")
            output_docx_path = os.path.join(OUTPUT_DIR, f"{base_name}_table.docx")

            # Generate JSON from PDF
            process_pdf(input_path, output_json_path)

            # Convert JSON to Word document
            json_to_word(output_json_path, output_docx_path)

if __name__ == "__main__":
    main()